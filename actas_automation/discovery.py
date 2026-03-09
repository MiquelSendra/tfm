"""Workspace discovery helpers."""

from __future__ import annotations

import logging
import re
from pathlib import Path

from docx import Document

from .models import SourceFiles
from .text_utils import normalize_text


def discover_source_files(workspace_dir: Path, logger: logging.Logger) -> SourceFiles:
    """Find the ZIP, Excel, template, and PDF files needed by the pipeline."""
    zip_file = _discover_zip_file(workspace_dir)
    excel_file = _discover_excel_file(workspace_dir)
    docx_template = _discover_docx_template(workspace_dir, logger)
    pdf_files = _discover_pdf_files(workspace_dir)

    logger.info("ZIP detected: %s", zip_file.name)
    logger.info("Excel detected: %s", excel_file.name)
    if docx_template:
        logger.info("DOCX template detected: %s", docx_template.name)
    else:
        logger.warning("No DOCX template detected in workspace.")
    logger.info("PDF files detected: %s", len(pdf_files))

    return SourceFiles(
        excel_file=excel_file,
        zip_file=zip_file,
        docx_template=docx_template,
        pdf_files=pdf_files,
    )


def _discover_zip_file(workspace_dir: Path) -> Path:
    zip_candidates = [
        path
        for path in workspace_dir.glob("*.zip")
        if re.fullmatch(r"\d+\.zip", path.name)
    ]
    if not zip_candidates:
        raise FileNotFoundError(
            "No ZIP file with numeric filename found (expected pattern: 123456.zip)."
        )
    zip_candidates.sort(key=lambda path: path.stat().st_mtime, reverse=True)
    return zip_candidates[0]


def _discover_excel_file(workspace_dir: Path) -> Path:
    excel_candidates = [
        path
        for path in workspace_dir.glob("*.xlsx")
        if not path.name.startswith("~$")
    ]
    if not excel_candidates:
        raise FileNotFoundError("No .xlsx file found in workspace.")
    excel_candidates.sort(key=lambda path: path.stat().st_mtime, reverse=True)
    return excel_candidates[0]


def _discover_docx_template(
    workspace_dir: Path, logger: logging.Logger
) -> Path | None:
    docx_candidates = sorted(workspace_dir.glob("*.docx"))
    if not docx_candidates:
        return None

    scored: list[tuple[int, Path]] = []
    for path in docx_candidates:
        try:
            text = _extract_docx_text(path)
            normalized = normalize_text(text)
            score = 0
            for marker in ("titulacion", "datos del estudiante", "tribunal"):
                if marker in normalized:
                    score += 1
            scored.append((score, path))
        except Exception as exc:  # pragma: no cover - best effort classification
            logger.warning("Could not inspect DOCX %s: %s", path.name, exc)
            scored.append((0, path))

    scored.sort(key=lambda item: item[0], reverse=True)
    return scored[0][1] if scored else None


def _extract_docx_text(path: Path) -> str:
    document = Document(path)
    chunks: list[str] = []
    chunks.extend(paragraph.text for paragraph in document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                chunks.append(cell.text)
    return "\n".join(chunks)


def _discover_pdf_files(workspace_dir: Path) -> list[Path]:
    pdf_files: list[Path] = []
    for path in workspace_dir.rglob("*.pdf"):
        if _should_skip_pdf(path, workspace_dir):
            continue
        pdf_files.append(path)
    return sorted(pdf_files)


def _should_skip_pdf(path: Path, workspace_dir: Path) -> bool:
    try:
        relative_parts = path.resolve().relative_to(workspace_dir.resolve()).parts
    except ValueError:
        return True

    if any(part.startswith(".") for part in relative_parts):
        return True
    if "output" in relative_parts:
        return True

    name = path.name
    if name.startswith("~$") or name.startswith("._"):
        return True
    return False
