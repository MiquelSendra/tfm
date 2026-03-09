"""DOCX generation from an acta template."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from .models import ActaContext
from .text_utils import normalize_text


def generate_docx_acta(template_path: Path, output_path: Path, context: ActaContext) -> None:
    """Create one acta DOCX by filling known sections in the template."""
    document = Document(str(template_path))

    filled_primary = _fill_expected_template_layout(document, context)
    if not filled_primary:
        _fill_using_label_fallback(document, context)

    document.save(str(output_path))


def _fill_expected_template_layout(document: Document, context: ActaContext) -> bool:
    """Fill the known template layout used in this process."""
    if len(document.tables) < 2:
        return False

    info_table = document.tables[0]
    tft_table = document.tables[1]

    if len(info_table.rows) < 6:
        return False
    if len(tft_table.rows) < 4:
        return False

    controls_written = [
        _set_first_content_control_text(info_table.cell(0, 0), context.titulacion),
        _set_first_content_control_text(info_table.cell(1, 0), context.edicion),
        _set_first_content_control_text(info_table.cell(3, 0), context.student_name),
        _set_first_content_control_text(info_table.cell(5, 0), context.dni),
        _set_first_content_control_text(tft_table.cell(1, 0), context.thesis_title),
        _set_first_content_control_text(tft_table.cell(3, 0), context.director),
    ]
    if not all(controls_written):
        return False

    _enforce_course_and_dni_fields(info_table, context)
    _set_cell_alignment(info_table.cell(5, 0), WD_ALIGN_PARAGRAPH.CENTER)
    return True


def _fill_using_label_fallback(document: Document, context: ActaContext) -> None:
    """Fallback filling strategy for templates with a different table layout."""
    mapping = {
        "titulacion": context.titulacion,
        "curso academico": context.edicion,
        "edicion": context.edicion,
        "nombre y apellidos": context.student_name,
        "dni pasaporte": context.dni,
        "titulo": context.thesis_title,
        "director": context.director,
    }

    for table in document.tables:
        for row in table.rows:
            for index, cell in enumerate(row.cells):
                normalized_cell = normalize_text(cell.text)
                for key, value in mapping.items():
                    if key in normalized_cell:
                        if index + 1 < len(row.cells):
                            row.cells[index + 1].text = value
                        else:
                            label = cell.text.strip().rstrip(":")
                            cell.text = f"{label}: {value}"
                        break


def _set_first_content_control_text(cell, value: str) -> bool:
    """Replace the first content control text in a cell."""
    sdts = cell._tc.xpath(".//*[local-name()='sdt']")
    if not sdts:
        return False

    sdt = sdts[0]
    for placeholder in sdt.xpath(
        "./*[local-name()='sdtPr']/*[local-name()='showingPlcHdr']"
    ):
        placeholder.getparent().remove(placeholder)

    text_nodes = sdt.xpath(
        ".//*[local-name()='sdtContent']//*[local-name()='t']"
    )
    if text_nodes:
        text_nodes[0].text = value
        for node in text_nodes[1:]:
            node.text = ""
        return True

    content = sdt.xpath(".//*[local-name()='sdtContent']")
    if not content:
        return False
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    paragraph.clear()
    paragraph.add_run(value)
    return True


def _set_cell_alignment(cell, alignment: WD_ALIGN_PARAGRAPH) -> None:
    """Set paragraph alignment for all cell paragraphs, including SDT blocks."""
    for paragraph in cell.paragraphs:
        paragraph.alignment = alignment

    for paragraph in cell._tc.xpath(
        ".//*[local-name()='sdtContent']//*[local-name()='p']"
    ):
        ppr = paragraph.find(qn("w:pPr"))
        if ppr is None:
            ppr = paragraph.makeelement(qn("w:pPr"))
            paragraph.insert(0, ppr)
        jc = ppr.find(qn("w:jc"))
        if jc is None:
            jc = paragraph.makeelement(qn("w:jc"))
            ppr.append(jc)
        jc.set(qn("w:val"), "center")


def _enforce_course_and_dni_fields(info_table, context: ActaContext) -> None:
    """Hard-enforce replacement of course and DNI placeholders."""
    if len(info_table.rows) < 6:
        return

    course_cell = info_table.cell(1, 0)
    _replace_text_token_in_cell(course_cell, "Curso académico", context.edicion)

    if len(info_table.rows[1].cells) > 1:
        # Avoid duplicated value in side cell from previous strategies.
        info_table.cell(1, 1).text = ""

    dni_cell = info_table.cell(5, 0)
    _set_cell_single_text(dni_cell, context.dni)


def _replace_text_token_in_cell(cell, placeholder: str, value: str) -> bool:
    """Replace exact token occurrences inside a cell XML tree."""
    replaced = False
    for text_node in cell._tc.xpath(".//*[local-name()='t']"):
        text_value = (text_node.text or "").strip()
        if text_value == placeholder:
            text_node.text = value
            replaced = True
    return replaced


def _set_cell_single_text(cell, value: str) -> None:
    """Keep only one visible text value in the cell."""
    text_nodes = cell._tc.xpath(".//*[local-name()='t']")
    if not text_nodes:
        cell.text = value
        return

    text_nodes[0].text = value
    for text_node in text_nodes[1:]:
        text_node.text = ""
